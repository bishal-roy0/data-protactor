"""Build the Karna accuracy evaluation report as a reviewed Word document."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Karna_Accuracy_Evaluation_Report.docx"
EVIDENCE = ROOT / "docs" / "evidence"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 100, 117)
RISK_RED = RGBColor(155, 28, 28)
BLACK = RGBColor(20, 29, 39)
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_font(run, size=11, color=BLACK, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(width_inches * 1440)))
    width.set(qn("w:type"), "dxa")


def style_cell(cell, *, header=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade(cell, HEADER_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_font(run, size=9, color=DARK_BLUE if header else BLACK, bold=header)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_body(doc, text, *, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    if bold_prefix and text.startswith(bold_prefix):
        set_font(paragraph.add_run(bold_prefix), bold=True)
        set_font(paragraph.add_run(text[len(bold_prefix):]))
    else:
        set_font(paragraph.add_run(text))
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph.add_run(text))
    return paragraph


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_cell_width(table.cell(0, 0), 6.5)
    cell = table.cell(0, 0)
    shade(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run(f"{label}: "), color=DARK_BLUE, bold=True)
    set_font(paragraph.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, (header, width) in enumerate(zip(headers, widths, strict=True)):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_width(cell, width)
        style_cell(cell, header=True)
    for row in rows:
        cells = table.add_row().cells
        for index, (value, width) in enumerate(zip(row, widths, strict=True)):
            cells[index].text = str(value)
            set_cell_width(cells[index], width)
            style_cell(cells[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def set_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run("Karna Accuracy Evaluation Report | "), size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_footer(section)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("KARNA | PRODUCTION AUDIT"), size=8, color=MUTED, bold=True)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def main():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("Karna Production Accuracy, Security, and Regression Audit"), size=23, color=BLACK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("Versioned v1 evaluation corpus | 28 July 2026"), size=12, color=MUTED)

    for label, value in (
        ("Scope", "Baseline text, URL, image-validation, dashboard, deployment, and failure-path review"),
        ("Dataset", "30 evaluated text/URL samples plus 4 image validation or capability records"),
        ("Result", "25 automated Python tests passed; Android tests were not run because this machine has no Java, Gradle, or Gradle wrapper"),
        ("Deployment", "https://data-protactor.vercel.app"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(f"{label}: "), bold=True)
        set_font(paragraph.add_run(value))

    add_heading(doc, "Executive conclusion")
    add_body(doc, "Karna passed 25 automated Python tests after this audit. The v1 corpus produced 20 exact category matches out of 30 evaluated text/URL samples. This is a diagnostic result from a small synthetic corpus, not a public real-world accuracy claim.")
    add_callout(doc, "Decision", "Keep Karna publicly described as an explainable, advisory threat-analysis API. Do not publish a numeric overall accuracy score from this corpus.")

    add_heading(doc, "Method and scope")
    add_body(doc, "Each corpus record has an ID, modality, expected category or behavior, and a non-sensitive sample. One-vs-rest metrics were calculated for each text/URL category. Image semantic scoring is not evaluated without a separately authorized OpenAI vision configuration.")
    for item in (
        "Safe look-alikes include benign urgency and a legitimate OTP-style message to expose false positives.",
        "Threat cases include phishing, social engineering, impersonation, prize/payment scams, suspicious URLs, and download links.",
        "The corpus includes obfuscated English, Spanish phishing, and Hindi social-engineering examples to expose language coverage limits.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Test execution")
    add_table(
        doc,
        ["Check", "Outcome", "Evidence"],
        [
            ("Python API and regression suite", "PASS", "25 passed"),
            ("Versioned v1 corpus", "COMPLETE", "30 evaluated text/URL samples"),
            ("Image validation", "PASS", "Unsupported type, size, and fallback coverage"),
            ("External failure paths", "PASS", "Simulated OpenAI and VirusTotal failures"),
            ("Live deployment", "PASS", "Health and URL-category regression"),
            ("Android unit suite", "NOT RUN", "No local Java, Gradle, or wrapper"),
        ],
        [2.1, 1.0, 3.4],
    )

    doc.add_page_break()
    add_heading(doc, "Category metrics")
    add_body(doc, "Metrics below are regression measurements for the v1 synthetic corpus only. They are not generalizable malware- or scam-detection accuracy figures.")
    add_table(
        doc,
        ["Category", "TP", "FP", "TN", "FN", "Precision", "Recall", "F1", "FPR"],
        [
            ("safe", 3, 6, 19, 2, "0.333", "0.600", "0.429", "0.240"),
            ("phishing", 5, 1, 20, 4, "0.833", "0.556", "0.667", "0.048"),
            ("social engineering", 3, 3, 22, 2, "0.500", "0.600", "0.545", "0.120"),
            ("impersonation", 2, 0, 27, 1, "1.000", "0.667", "0.800", "0.000"),
            ("scam", 2, 0, 27, 1, "1.000", "0.667", "0.800", "0.000"),
            ("suspicious URL", 2, 0, 28, 0, "1.000", "1.000", "1.000", "0.000"),
            ("malware download", 3, 0, 27, 0, "1.000", "1.000", "1.000", "0.000"),
        ],
        [1.35, 0.42, 0.42, 0.42, 0.42, 0.78, 0.64, 0.5, 0.5],
    )
    add_heading(doc, "Known limitations")
    for item in (
        "Benign urgency and legitimate OTP-style messages can receive cautionary classifications.",
        "Obfuscated phishing, less explicit scams, and non-English examples are currently missed or classified too broadly.",
        "The baseline never opens, downloads, crawls, or executes URLs; therefore it cannot claim a complete malware scan.",
        "Image semantics are not measured without an authorized OpenAI vision key and a separately labeled image corpus.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Audit findings and fixes")
    add_callout(doc, "Fixed", "Deceptive media-link evidence now returns suspicious_url instead of social_engineering.")
    add_callout(doc, "Fixed", "Prize, lottery, and processing-fee cues now activate the scam category; a regression test protects the behavior.")
    add_body(doc, "The unresolved false positives and false negatives remain in the corpus. They should drive future multilingual and semantic-analysis work rather than being removed from the audit.")

    doc.add_page_break()
    add_heading(doc, "Security and resilience checks")
    for item in (
        "Submitted URLs are structurally analyzed only; the baseline does not visit, download, crawl, or execute them.",
        "A simulated VirusTotal failure returns no external evidence and preserves the local result.",
        "A simulated OpenAI vision failure returns a structured caution response without leaking the provider error.",
        "Unsupported image types and images larger than 5 MB are rejected, and uploads are not persistently stored.",
        "The public configuration endpoint exposes only the Android download URL and never API keys.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Live test evidence")
    add_body(doc, "The following screenshots show live dashboard checks. The phishing case returned critical/phishing/quarantine. The executable URL case returned high/malware download/block.")
    for title, filename in (
        ("Live phishing result", "phishing-dashboard-result.png"),
        ("Live executable-download URL result", "malware-url-dashboard-result.png"),
    ):
        caption = doc.add_paragraph()
        caption.paragraph_format.space_before = Pt(8)
        caption.paragraph_format.space_after = Pt(4)
        set_font(caption.add_run(title), size=11, color=DARK_BLUE, bold=True)
        doc.add_picture(str(EVIDENCE / filename), width=Inches(6.2))

    add_heading(doc, "Reproducibility and next steps")
    add_body(doc, "Run the versioned evaluation with: PYTHONPATH=src, then python scripts/run_baseline_evaluation.py and python -m pytest -q. The machine-readable output is docs/evidence/evaluation_metrics_v1.json.")
    for item in (
        "Build a larger consented or lawfully sourced corpus with independent labels.",
        "Use development and holdout splits; record label-review process and sample provenance.",
        "Add multilingual, adversarial, and score-calibration evaluation before publishing any production metric.",
        "Run Android unit tests on a configured Android Studio or CI runner before an Android release.",
    ):
        add_bullet(doc, item)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
